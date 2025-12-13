# tournament_manager.py
import json
import socket
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QFrame, QGridLayout, QGroupBox,
    QLineEdit, QTextEdit, QMessageBox, QFileDialog, QTabWidget,
    QTableWidget, QTableWidgetItem, QComboBox, QListWidget,
    QSplitter, QProgressBar, QHeaderView, QDialog, QDialogButtonBox, QFormLayout, QDesktopWidget
)
from PyQt5.QtCore import Qt, QTimer, pyqtSignal
from PyQt5.QtGui import QFont, QScreen, QPainter, QPen, QBrush, QColor, QPixmap
from core.utils import create_bracket, generate_schedule, get_wrestler_club
from core.settings import get_settings
from ui.dialogs.wrestler_dialogs import AddWrestlerDialog, MoveWrestlerDialog
from ui.dialogs.category_dialogs import CategoryEditDialog
from PyQt5.QtWidgets import QGraphicsView, QGraphicsScene, QGraphicsRectItem, QGraphicsTextItem, QColorDialog
from PyQt5.QtNetwork import QTcpServer, QTcpSocket, QHostAddress
from ui.widgets.schedule import MatScheduleWindow


class BracketWidget(QGraphicsView):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.scene = QGraphicsScene()
        self.setScene(self.scene)
        self.setRenderHint(QPainter.Antialiasing)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.matches = []
        self.match_height = 30  # Уменьшили высоту матча
        self.match_width = 150  # Уменьшили ширину матча
        self.vertical_spacing = 3  # Уменьшили вертикальный отступ
    
    def set_matches(self, matches):
        self.matches = matches
        self.draw_bracket()
    
    def draw_bracket(self):
        self.scene.clear()
        
        if not self.matches:
            return
        
        # Группируем матчи по раундам для отображения заголовков
        rounds = {}
        for match in self.matches:
            round_num = match.get('round', 1)
            if round_num not in rounds:
                rounds[round_num] = []
            rounds[round_num].append(match)
        
        # Вычисляем оптимальное расположение
        x = 10  # Фиксированная позиция X для всех матчей
        y = 10  # Начальная позиция Y
        
        # Сортируем раунды для правильного порядка отображения
        for round_num in sorted(rounds.keys()):
            matches = rounds[round_num]
                        
            # Отображаем все матчи раунда вертикально
            for match in matches:
                rect = self.scene.addRect(x, y, self.match_width, self.match_height, 
                                         QPen(Qt.black, 1))  # Уменьшили толщину линии
                
                if match.get('completed'):
                    rect.setBrush(QBrush(QColor("#9ba6bd")))
                else:
                    rect.setBrush(QBrush(Qt.white))
                
                rect.match_data = match
                
                w1 = match.get('wrestler1', '')
                w2 = match.get('wrestler2', '')
                score = f"{match.get('score1', 0)}:{match.get('score2', 0)}"
                winner = match.get('winner', '')
                
                # Уменьшаем размер текста
                text1 = QGraphicsTextItem(f"◉ {w1}")
                text1.setFont(QFont("Arial", 7))
                text1.setPos(x + 2, y + 2)  # Уменьшили отступ
                self.scene.addItem(text1)
                
                text2 = QGraphicsTextItem(f"◉ {w2}")
                text2.setFont(QFont("Arial", 7))
                text2.setPos(x + 2, y + self.match_height//2 + 1)
                self.scene.addItem(text2)
                
                score_text = QGraphicsTextItem(score)
                score_text.setFont(QFont("Arial", 7))
                score_text.setPos(x + self.match_width - 25, y + self.match_height//2 - 5)
                self.scene.addItem(score_text)
                
                if winner:
                    if winner == w1:
                        text1.setDefaultTextColor(QColor(0, 100, 0))
                        text1.setFont(QFont("Arial", 7, QFont.Bold))
                    elif winner == w2:
                        text2.setDefaultTextColor(QColor(0, 100, 0))
                        text2.setFont(QFont("Arial", 7, QFont.Bold))
                
                y += self.match_height + self.vertical_spacing
            
            # Дополнительный отступ между раундами
            if len(rounds) > 1:
                y += 3  # Уменьшили отступ между раундами
        
        # Добавляем небольшие отступы по краям
        scene_rect = self.scene.itemsBoundingRect()
        margin = 5
        scene_rect.adjust(-margin, -margin, margin, margin)
        self.scene.setSceneRect(scene_rect)
        
        # Автоматически масштабируем вид, чтобы вся сцена поместилась
        # Используем KeepAspectRatio, чтобы не растягивать элементы
        self.fitInView(self.scene.sceneRect(), Qt.KeepAspectRatio)
    
    def resizeEvent(self, event):
        """Обработка изменения размера - перемасштабирование"""
        super().resizeEvent(event)
        if self.scene and self.scene.items():
            # Используем KeepAspectRatio, чтобы не растягивать элементы
            self.fitInView(self.scene.sceneRect(), Qt.KeepAspectRatio)
    
    def mousePressEvent(self, event):
        """Обработка клика по матчу в графическом представлении"""
        pos = self.mapToScene(event.pos())
        items = self.scene.items(pos)
        
        for item in items:
            if hasattr(item, 'match_data'):
                parent = self.parent()
                while parent and not isinstance(parent, BracketWindow):
                    parent = parent.parent()
                
                if parent and hasattr(parent, 'load_match_from_bracket'):
                    parent.load_match_from_bracket(item.match_data)
                break
        
        super().mousePressEvent(event)



class BracketWindow(QMainWindow):
    match_autoload = pyqtSignal(dict)
    
    def __init__(self, parent=None, tournament_data=None):
        super().__init__(parent)
        # Локальная копия нужна только как запасной вариант,
        # актуальные данные всегда берём из главного окна.
        self.tournament_data = tournament_data
        self.setWindowTitle("Табличное и графическое представление сетки")
        self.current_category = None
        self._elim_row_to_match = {}
        self._round_robin_participants = []
        self.setup_ui()
        
        print("[DEBUG] BracketWindow создан, сигнал match_autoload:", hasattr(self, 'match_autoload'))

    def setup_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # Кнопка закрытия окна
        close_btn = QPushButton("✕ Закрыть")
        close_btn.clicked.connect(self.close)
        close_btn.setStyleSheet("background-color: #c62828; color: white; font-weight: bold; padding: 5px;")
        layout.addWidget(close_btn)

        # Используем QSplitter для возможности изменения размера областей
        splitter = QSplitter(Qt.Vertical)
        splitter.setChildrenCollapsible(False)
        layout.addWidget(splitter)

        # === Круговая таблица категории ===
        # Используем QWidget вместо QGroupBox для полного контроля отступов
        self.round_group = QWidget()
        round_layout = QVBoxLayout(self.round_group)
        # Устанавливаем отступы: верхний, левый, правый, нижний (последний = 0)
        round_layout.setContentsMargins(0, 0, 0, 0)
        round_layout.setSpacing(0)
        
        # Добавляем заголовок
        round_title = QLabel("Табличное представление")
        round_title.setStyleSheet("font-weight: bold; padding: 5px;")
        round_layout.addWidget(round_title)
        
        self.round_table = QTableWidget()
        self.round_table.setEditTriggers(QTableWidget.NoEditTriggers)
        # ПРАВИЛЬНОЕ ПОДКЛЮЧЕНИЕ ДВОЙНОГО КЛИКА
        self.round_table.cellDoubleClicked.connect(self.on_round_table_double_click)
        # Настраиваем политику размера: таблица не должна растягиваться по вертикали
        from PyQt5.QtWidgets import QSizePolicy
        self.round_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.round_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        self.round_table.verticalHeader().setVisible(False)
        # Убираем все отступы у таблицы
        self.round_table.setStyleSheet("QTableWidget { margin: 0px; padding: 0px; }")
        round_layout.addWidget(self.round_table)
        splitter.addWidget(self.round_group)

        # Графическое представление сетки (слева под таблицей)
        self.bracket_group = QGroupBox("Графическая сетка")
        # Настраиваем политику размера: графическая сетка должна растягиваться
        from PyQt5.QtWidgets import QSizePolicy
        self.bracket_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        bracket_main_layout = QVBoxLayout(self.bracket_group)
        bracket_h_layout = QHBoxLayout()
        # Выравнивание содержимого слева
        bracket_h_layout.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self.bracket_widget = BracketWidget()
        # Настраиваем политику размера для виджета сетки
        self.bracket_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        bracket_h_layout.addWidget(self.bracket_widget)
        # Убрали отступ справа и добавили растягивающийся элемент
        bracket_h_layout.addStretch()  # Это займет все свободное место справа
        bracket_main_layout.addLayout(bracket_h_layout)
        splitter.addWidget(self.bracket_group)
        # Устанавливаем stretch factor: таблица не растягивается (0), графическая сетка растягивается (1)
        # Это позволяет графической сетке занимать все оставшееся место
        splitter.setStretchFactor(0, 0)  # Таблица - не растягивается
        splitter.setStretchFactor(1, 1)  # Графическая сетка - растягивается

        buttons_layout = QHBoxLayout()
        broadcast_btn = QPushButton("Транслировать на второй экран")
        broadcast_btn.clicked.connect(self.broadcast_to_second_screen)
        buttons_layout.addWidget(broadcast_btn)
        
        export_btn = QPushButton("📄 Экспортировать в DOCX")
        export_btn.clicked.connect(self.export_to_docx)
        export_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold; padding: 5px;")
        buttons_layout.addWidget(export_btn)
        
        layout.addLayout(buttons_layout)

    def find_control_panel_by_mat(self, mat_number):
        """Находит панель управления по номеру ковра"""
        # Ищем среди всех виджетов панель управления с заданным номером ковра
        for widget in QApplication.allWidgets():
            if hasattr(widget, 'mat_number') and widget.mat_number == mat_number:
                return widget
        
        # Если не нашли, ищем по заголовку вкладки
        main_window = self.get_main_window()
        if main_window and hasattr(main_window, 'tab_widget'):
            for i in range(main_window.tab_widget.count()):
                if main_window.tab_widget.tabText(i) == f"Управление — Ковёр {mat_number}":
                    return main_window.tab_widget.widget(i)
        
        return None

    def get_main_window(self):
        """Находит главное окно приложения"""
        # Ищем среди всех top-level виджетов главное окно
        for widget in QApplication.topLevelWidgets():
            # Проверяем, что это не текущее окно (BracketWindow) и имеет нужные методы
            if widget != self and hasattr(widget, 'open_control_panel_tab'):
                return widget
        return None

    def load_match_from_bracket(self, match_data):
        """Загружает матч из графического представления в панель управления"""
        # Всегда берём актуальные данные турнира из главного окна
        main_window = self.get_main_window()
        tournament_data = getattr(main_window, 'tournament_data', None) if main_window else self.tournament_data
        if not tournament_data:
            return

        if not main_window:
            QMessageBox.warning(self, "Ошибка", "Не удалось найти главное окно приложения")
            return

        w1 = match_data.get('wrestler1', '')
        w2 = match_data.get('wrestler2', '')

        if not w1 or not w2:
            return

        # Ищем клуб в категории, если не найден в глобальном списке
        def find_club(wrestler_name):
            club = get_wrestler_club(tournament_data, wrestler_name)
            if not club and self.current_category:
                category = tournament_data.get('categories', {}).get(self.current_category, {})
                participants = category.get('participants', []) or category.get('wrestlers', [])
                for p in participants:
                    if p.get('name') == wrestler_name:
                        club = p.get('club', '') or p.get('region', '') or p.get('тренер', '')
                        break
            return club
        
        w1_data = {
            'name': w1,
            'club': find_club(w1),
            'category': self.current_category
        }
        w2_data = {
            'name': w2,
            'club': find_club(w2),
            'category': self.current_category
        }

        # Проверяем, не завершен ли уже матч
        if match_data.get('completed'):
            reply = QMessageBox.question(self, 'Загрузка матча',
                                        'Этот матч уже завершен. Загрузить его для редактирования?',
                                        QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.No:
                return

        # ОБНОВЛЕНО: Убедимся, что панель управления существует
        cp = self.find_control_panel_by_mat(1)
        if not cp:
            # Если панель не найдена, создаем ее
            if hasattr(main_window, 'open_control_panel_tab'):
                main_window.open_control_panel_tab(mat_number=1)
                # Даем время на создание
                QApplication.processEvents()
                cp = self.find_control_panel_by_mat(1)

        if cp:
            cp.set_match_competitors(w1_data, w2_data)

            # Если матч уже сыгран, загружаем счет
            if match_data.get('completed'):
                cp.red.points = match_data.get('score1', 0)
                cp.blue.points = match_data.get('score2', 0)
                cp.update_display()

            if hasattr(cp, 'set_current_match_info'):
                cp.set_current_match_info(self.current_category, w1, w2, match_data.get('id'))

            cp.send_scoreboard_update()
            # Дополнительная отправка через 100 мс
            QTimer.singleShot(100, cp.send_scoreboard_update)
            QMessageBox.information(self, "Загрузка матча", f"Матч загружен:\n{w1} vs {w2}")
        else:
            QMessageBox.warning(self, "Ошибка", "Не удалось найти или создать панель управления")

    def find_next_unplayed_match(self):
        """Возвращает первый несыгранный матч текущей категории"""
        main_window = self.get_main_window()
        tournament_data = getattr(main_window, 'tournament_data', None) if main_window else self.tournament_data
        if not tournament_data or not self.current_category:
            return None
    
        category = tournament_data["categories"].get(self.current_category)
        if not category:
            return None
    
        matches = category.get("matches", [])
        for m in matches:
            if not m.get("completed", False):
                return m
    
        return None
    
    def update_bracket(self, cat):
        main_window = self.get_main_window()
        tournament_data = getattr(main_window, 'tournament_data', None) if main_window else self.tournament_data
        if not tournament_data:
            return

        self.current_category = cat

        category = tournament_data['categories'].get(cat)
        if not category:
            return

        matches = category.get('matches', [])

        self.bracket_widget.set_matches(matches)
        self.update_round_robin_table(cat)
        # Автозагрузка только если нет активного матча
        # autoload_match теперь сам проверяет наличие активного матча
        self.autoload_match()

    def autoload_match(self):
        """Автозагрузка следующего несыгранного матча в панель управления"""
        print(f"[DEBUG] autoload_match вызван для категории: {self.current_category}")

        # Проверяем, есть ли уже активный матч в панели управления
        main_window = self.get_main_window()
        if main_window:
            cp = main_window.find_control_panel_by_mat(1)
            if cp:
                # Проверяем, есть ли уже загруженный матч
                if hasattr(cp, 'current_match_w1') and hasattr(cp, 'current_match_w2'):
                    if cp.current_match_w1 and cp.current_match_w2:
                        # Проверяем, не завершен ли текущий матч
                        if hasattr(cp, 'current_match_id') and cp.current_match_id:
                            # Ищем матч в категории
                            category = self.tournament_data.get('categories', {}).get(self.current_category, {})
                            matches = category.get('matches', [])
                            for m in matches:
                                if m.get('id') == cp.current_match_id:
                                    # Если матч не завершен, не заменяем его
                                    if not m.get('completed', False):
                                        print(f"[DEBUG] Пропускаем автозагрузку: уже есть активный матч {cp.current_match_w1} vs {cp.current_match_w2}")
                                        return
                                    break

        match = self.find_next_unplayed_match()
        if not match:
            print("[DEBUG] Все матчи сыграны или матч не найден")
            return

        w1 = match.get("wrestler1", "")
        w2 = match.get("wrestler2", "")

        print(f"[DEBUG] Найден несыгранный матч: {w1} vs {w2}")

        data = {
            "category": self.current_category,
            "w1": w1,
            "w2": w2,
            "match_id": match.get("id"),
            "score1": match.get("score1", 0),
            "score2": match.get("score2", 0),
            "completed": match.get("completed", False),
        }

        print(f"[DEBUG] Отправка данных через сигнал: {data}")
        self.match_autoload.emit(data)

    def update_round_robin_table(self, cat):
        """Обновляет таблицу круговой системы"""
        main_window = self.get_main_window()
        tournament_data = getattr(main_window, 'tournament_data', None) if main_window else self.tournament_data
        if not cat or not tournament_data:
            self.round_group.setVisible(False)
            return
        
        # Сохраняем текущую позицию прокрутки
        scroll_pos = self.round_table.verticalScrollBar().value()
        h_scroll_pos = self.round_table.horizontalScrollBar().value()
        
        self.round_table.clear()

        category = tournament_data['categories'].get(cat)
        if not category:
            self.round_group.setVisible(False)
            return

        cat_type = category.get('type', 'round_robin')

        if cat_type != 'round_robin':
            self._update_elimination_table(category, cat)
            return

        participants = [p.get('name', '') for p in category.get('participants', [])]
        participants = [name for name in participants if name]
        self._round_robin_participants = participants

        n = len(participants)
        if n == 0:
            self.round_group.setVisible(False)
            return

        self.round_group.setVisible(True)

        matches = category.get('matches', [])
        index_by_name = {name: idx for idx, name in enumerate(participants)}

        results = [["" for _ in range(n)] for _ in range(n)]
        stats = {name: {"wins": 0, "losses": 0, "points": 0} for name in participants}

        for m in matches:
            w1 = m.get('wrestler1')
            w2 = m.get('wrestler2')
            if w1 not in index_by_name or w2 not in index_by_name:
                continue

            i = index_by_name[w1]
            j = index_by_name[w2]

            s1 = int(m.get('score1', 0) or 0)
            s2 = int(m.get('score2', 0) or 0)
            completed = m.get('completed', False)
            winner = m.get('winner')

            if completed:
                # В круговой таблице показываем не реальный счёт (0:10),
                # а "очко за победу": 1 победителю, 0 проигравшему.
                if s1 > s2 or winner == w1:
                    results[i][j] = "1"
                    results[j][i] = "0"
                elif s2 > s1 or winner == w2:
                    results[i][j] = "0"
                    results[j][i] = "1"
                else:
                    # Ничья — по 0
                    results[i][j] = "0"
                    results[j][i] = "0"

                if winner == w1 or s1 > s2:
                    stats[w1]["wins"] += 1
                    stats[w2]["losses"] += 1
                    stats[w1]["points"] += 1
                elif winner == w2 or s2 > s1:
                    stats[w2]["wins"] += 1
                    stats[w1]["losses"] += 1
                    stats[w2]["points"] += 1
            else:
                results[i][j] = ""
                results[j][i] = ""

        # Вычисляем места участников
        def calculate_place(name1, name2):
            """Сравнивает двух участников для определения места"""
            stats1 = stats.get(name1, {"wins": 0, "losses": 0, "points": 0})
            stats2 = stats.get(name2, {"wins": 0, "losses": 0, "points": 0})
            
            # Сначала по количеству побед
            if stats1["wins"] > stats2["wins"]:
                return -1
            elif stats1["wins"] < stats2["wins"]:
                return 1
            
            # Если одинаковое количество побед, проверяем личную встречу
            i1 = index_by_name.get(name1)
            i2 = index_by_name.get(name2)
            if i1 is not None and i2 is not None:
                head_to_head_1 = results[i1][i2]  # Результат name1 против name2
                head_to_head_2 = results[i2][i1]  # Результат name2 против name1
                if head_to_head_1 == "1":
                    return -1  # name1 победил name2
                elif head_to_head_2 == "1":
                    return 1   # name2 победил name1
            
            # Если личная встреча не состоялась или ничья, сортируем по очкам
            if stats1["points"] > stats2["points"]:
                return -1
            elif stats1["points"] < stats2["points"]:
                return 1
            
            # В крайнем случае - по алфавиту
            return -1 if name1 < name2 else 1
        
        # Сортируем участников по местам
        from functools import cmp_to_key
        sorted_participants = sorted(participants, key=cmp_to_key(calculate_place))
        
        # Вычисляем места (с учетом одинаковых результатов)
        places = {}
        current_place = 1
        for idx, name in enumerate(sorted_participants):
            if idx > 0:
                prev_name = sorted_participants[idx - 1]
                prev_stats = stats.get(prev_name, {"wins": 0, "points": 0})
                curr_stats = stats.get(name, {"wins": 0, "points": 0})
                
                # Проверяем, нужно ли увеличить место
                if prev_stats["wins"] != curr_stats["wins"]:
                    # Разное количество побед - новое место
                    current_place = idx + 1
                else:
                    # Одинаковое количество побед - проверяем личную встречу
                    i_prev = index_by_name.get(prev_name)
                    i_curr = index_by_name.get(name)
                    should_increment = False
                    
                    if i_prev is not None and i_curr is not None:
                        h2h_prev = results[i_prev][i_curr]
                        h2h_curr = results[i_curr][i_prev]
                        # Если есть результат личной встречи и один победил другого
                        if h2h_prev == "1" or h2h_curr == "1":
                            # Есть результат личной встречи - новое место
                            should_increment = True
                        elif prev_stats["points"] != curr_stats["points"]:
                            # Разные очки - новое место
                            should_increment = True
                    
                    if should_increment:
                        current_place = idx + 1
                    # Иначе остаемся на том же месте (одинаковое место)
            
            places[name] = current_place
        
        columns = 1 + 1 + n + 3  # Место + Участник + участники + Победы + Поражения + Очки
        self.round_table.setRowCount(n)
        self.round_table.setColumnCount(columns)

        headers = ["Место", "Участник"]
        headers.extend(participants)
        headers.extend(["Победы", "Поражения", "Очки"])
        self.round_table.setHorizontalHeaderLabels(headers)

        # Отображаем в отсортированном порядке
        for row, name in enumerate(sorted_participants):
            # Место
            place_item = QTableWidgetItem(str(places[name]))
            place_item.setTextAlignment(Qt.AlignCenter)
            place_font = QFont()
            place_font.setBold(True)
            place_item.setFont(place_font)
            self.round_table.setItem(row, 0, place_item)
            
            # Участник
            name_item = QTableWidgetItem(name)
            self.round_table.setItem(row, 1, name_item)

            # Результаты встреч
            orig_row = index_by_name[name]
            for col, opp_name in enumerate(participants):
                table_col = 2 + col
                if orig_row == col:
                    cell = QTableWidgetItem("—")
                    cell.setTextAlignment(Qt.AlignCenter)
                else:
                    value = results[orig_row][col]
                    cell = QTableWidgetItem(value)
                    cell.setTextAlignment(Qt.AlignCenter)
                    if value:
                        cell.setBackground(QColor("#9ba6bd"))
                self.round_table.setItem(row, table_col, cell)

            st = stats.get(name, {"wins": 0, "losses": 0, "points": 0})
            wins_item = QTableWidgetItem(str(st["wins"]))
            losses_item = QTableWidgetItem(str(st["losses"]))
            points_item = QTableWidgetItem(str(st["points"]))
            for item in (wins_item, losses_item, points_item):
                item.setTextAlignment(Qt.AlignCenter)

            self.round_table.setItem(row, 2 + n, wins_item)
            self.round_table.setItem(row, 2 + n + 1, losses_item)
            self.round_table.setItem(row, 2 + n + 2, points_item)
        
        # Восстанавливаем позицию прокрутки
        self.round_table.verticalScrollBar().setValue(scroll_pos)
        self.round_table.horizontalScrollBar().setValue(h_scroll_pos)
        
        # Устанавливаем высоту таблицы по содержимому
        self._adjust_table_height()

    def _adjust_table_height(self):
        """Устанавливает минимальную высоту таблицы по её содержимому"""
        if self.round_table.rowCount() == 0:
            return
        
        # Высота заголовка
        header_height = self.round_table.horizontalHeader().height()
        
        # Высота всех строк
        total_row_height = sum(self.round_table.rowHeight(i) for i in range(self.round_table.rowCount()))
        
        # Общая высота таблицы (заголовок + строки + границы)
        total_height = header_height + total_row_height + 2  # +2 для границ
        
        # Устанавливаем минимальную высоту таблицы по содержимому (не фиксированную)
        # Это позволит графической сетке занимать оставшееся место
        self.round_table.setMinimumHeight(total_height)

    def _update_elimination_table(self, category, cat_name):
        matches = category.get('matches', [])
        if not matches:
            self.round_group.setVisible(False)
            return

        self.round_group.setVisible(True)
        self._elim_row_to_match = {}

        rounds = {}
        for match in matches:
            round_num = match.get('round', 1)
            if round_num not in rounds:
                rounds[round_num] = []
            rounds[round_num].append(match)

        max_round = max(rounds.keys()) if rounds else 1

        round_names = {}
        for i in range(1, max_round + 1):
            if i == max_round:
                round_names[i] = "ФИНАЛ"
            elif i == max_round - 1:
                round_names[i] = "ПОЛУФИНАЛ"
            elif i == max_round - 2:
                round_names[i] = "ЧЕТВЕРТЬФИНАЛ"
            else:
                round_names[i] = f"РАУНД {i}"

        total_matches = len(matches)

        headers = ["Раунд", "Красный", "Синий", "Счет", "Победитель"]
        self.round_table.setColumnCount(len(headers))
        self.round_table.setHorizontalHeaderLabels(headers)
        self.round_table.setRowCount(total_matches)

        row = 0
        for round_num in sorted(rounds.keys()):
            round_matches = rounds[round_num]

            if len(round_matches) > 0:
                round_item = QTableWidgetItem(round_names[round_num])
                round_item.setBackground(QColor(240, 240, 240))
                round_item.setTextAlignment(Qt.AlignCenter)
                self.round_table.setItem(row, 0, round_item)
                self.round_table.setSpan(row, 0, 1, 5)
                row += 1

            for match in round_matches:
                round_item = QTableWidgetItem(str(round_num))
                round_item.setTextAlignment(Qt.AlignCenter)
                self.round_table.setItem(row, 0, round_item)

                wrestler1_item = QTableWidgetItem(match.get('wrestler1', ''))
                wrestler2_item = QTableWidgetItem(match.get('wrestler2', ''))

                score1 = match.get('score1', 0)
                score2 = match.get('score2', 0)
                score_item = QTableWidgetItem(f"{score1} : {score2}")
                score_item.setTextAlignment(Qt.AlignCenter)

                winner = match.get('winner', '')
                winner_item = QTableWidgetItem(winner)

                if winner:
                    if winner == match.get('wrestler1'):
                        wrestler1_item.setBackground(QColor(255, 255, 200))
                        wrestler1_item.setForeground(QBrush(QColor(0, 100, 0)))
                    elif winner == match.get('wrestler2'):
                        wrestler2_item.setBackground(QColor(255, 255, 200))
                        wrestler2_item.setForeground(QBrush(QColor(0, 100, 0)))

                if match.get('completed'):
                    for col_idx in range(5):
                        item = self.round_table.item(row, col_idx)
                        if item is None:
                            item = QTableWidgetItem("")
                            self.round_table.setItem(row, col_idx, item)
                        item.setBackground(QColor("#9ba6bd"))

                self.round_table.setItem(row, 1, wrestler1_item)
                self.round_table.setItem(row, 2, wrestler2_item)
                self.round_table.setItem(row, 3, score_item)
                self.round_table.setItem(row, 4, winner_item)

                self._elim_row_to_match[row] = match

                row += 1

        self.round_table.verticalHeader().setVisible(False)
        self.round_table.horizontalHeader().setStretchLastSection(True)

        self.round_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.round_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.round_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.round_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.round_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.Stretch)
        
        # Устанавливаем высоту таблицы по содержимому
        self._adjust_table_height()
    
    def on_round_table_double_click(self, row, column):
        """Двойной клик в таблице сетки – автоустановка борцов в панель управления схваткой."""
        print(f"[DEBUG] Двойной клик по ячейке: строка={row}, столбец={column}")

        if not self.tournament_data or not self.current_category:
            return

        category = self.tournament_data['categories'].get(self.current_category)
        if not category:
            return

        cat_type = category.get('type', 'round_robin')

        wrestler1_name = None
        wrestler2_name = None
        match_data = None

        if cat_type == 'round_robin':
            participants = self._round_robin_participants or []
            n = len(participants)
            if n == 0:
                return

            # Структура таблицы: столбец 0 = Место, столбец 1 = Участник, столбцы 2..(2+n-1) = результаты встреч
            # Клик должен быть по столбцам результатов встреч (2..(2+n-1))
            if column < 2 or column >= (2 + n):
                return

            if row < 0 or row >= n:
                return

            # Имя первого участника берем из столбца 1 (Участник), а не из столбца 0 (Место)
            name_item_row = self.round_table.item(row, 1)
            if not name_item_row:
                return
            wrestler1_name = name_item_row.text()

            # Индекс противника в списке participants: столбец - 2
            opp_index = column - 2
            if opp_index < 0 or opp_index >= n:
                return

            # Проверяем, что это не та же строка (диагональ таблицы)
            # Нужно проверить, что row в отсортированном списке не равен opp_index в исходном списке
            # Для этого найдем имя участника в строке row и его индекс в исходном списке
            wrestler1_orig_index = participants.index(wrestler1_name) if wrestler1_name in participants else -1
            if wrestler1_orig_index == opp_index:
                return

            # Имя второго участника берем из заголовка столбца или из списка participants
            header_item = self.round_table.horizontalHeaderItem(column)
            if header_item:
                wrestler2_name = header_item.text()
            else:
                if opp_index < len(participants):
                    wrestler2_name = participants[opp_index]
                else:
                    return

            for m in category.get('matches', []):
                w1 = m.get('wrestler1', '')
                w2 = m.get('wrestler2', '')
                if (w1 == wrestler1_name and w2 == wrestler2_name) or (w1 == wrestler2_name and w2 == wrestler1_name):
                    match_data = m
                    break
        else:
            match_data = self._elim_row_to_match.get(row)
            if not match_data:
                return
            wrestler1_name = match_data.get('wrestler1', '')
            wrestler2_name = match_data.get('wrestler2', '')

        if not wrestler1_name or not wrestler2_name:
            QMessageBox.warning(self, "Ошибка", "Не удалось определить участников матча")
            return

        # Получаем главное окно приложения
        app = QApplication.instance()
        main_window = None
        for widget in app.topLevelWidgets():
            # Ищем главное окно (не текущее BracketWindow)
            if widget != self and hasattr(widget, 'open_control_panel_tab'):
                main_window = widget
                break

        if not main_window:
            QMessageBox.warning(self, "Ошибка", "Не удалось найти главное окно приложения")
            return

        # Формируем данные для панели управления
        w1_data = {
            'name': wrestler1_name,
            'club': get_wrestler_club(self.tournament_data, wrestler1_name),
            'category': self.current_category
        }
        w2_data = {
            'name': wrestler2_name,
            'club': get_wrestler_club(self.tournament_data, wrestler2_name),
            'category': self.current_category
        }

        # Открываем панель управления
        if hasattr(main_window, 'open_control_panel_tab'):
            main_window.open_control_panel_tab(mat_number=1)

        # Ищем панель управления
        cp = None
        if hasattr(main_window, 'find_control_panel_by_mat'):
            cp = main_window.find_control_panel_by_mat(1)
        else:
            # Ручной поиск
            for widget in app.topLevelWidgets():
                if hasattr(widget, 'mat_number') and widget.mat_number == 1:
                    cp = widget
                    break

        if cp:
            cp.set_match_competitors(w1_data, w2_data)

            # Если матч уже сыгран, загружаем счет
            if match_data and match_data.get('completed'):
                # Определяем, кто красный, кто синий
                if match_data.get('wrestler1') == wrestler1_name:
                    cp.red.points = match_data.get('score1', 0)
                    cp.blue.points = match_data.get('score2', 0)
                else:
                    cp.red.points = match_data.get('score2', 0)
                    cp.blue.points = match_data.get('score1', 0)

                cp.update_display()

            if hasattr(cp, 'set_current_match_info'):
                match_id = match_data.get('id') if match_data else None
                cp.set_current_match_info(self.current_category, wrestler1_name, wrestler2_name, match_id)

            cp.send_scoreboard_update()
            QMessageBox.information(self, "Загрузка матча", f"Матч загружен:\n{w1_data['name']} vs {w2_data['name']}")
        else:
            QMessageBox.warning(self, "Ошибка", "Не удалось найти панель управления")

    def broadcast_to_second_screen(self):
        screens = QApplication.screens()
        if len(screens) < 2:
            QMessageBox.warning(self, "Ошибка", "Второй экран не обнаружен")
            return

        second_screen = screens[1]
        self.move_to_screen(second_screen)

    def move_to_screen(self, screen):
        geometry = screen.geometry()
        self.move(geometry.left(), geometry.top())
        self.showFullScreen()

    def export_to_docx(self):
        """Экспортирует страницу с сетками в DOCX файл"""
        import tempfile
        import os
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.warning(
                self, 
                "Ошибка", 
                "Библиотека python-docx не установлена.\n\n"
                "Установите её командой:\npip install python-docx"
            )
            return

        # Получаем данные турнира
        main_window = self.get_main_window()
        tournament_data = getattr(main_window, 'tournament_data', None) if main_window else self.tournament_data
        
        if not tournament_data or not tournament_data.get('categories'):
            QMessageBox.warning(self, "Ошибка", "Нет данных турнира для экспорта")
            return
        
        categories = list(tournament_data.get('categories', {}).keys())
        
        if not categories:
            QMessageBox.warning(self, "Ошибка", "Нет категорий для экспорта")
            return
        
        # Спрашиваем, что экспортировать
        from PyQt5.QtWidgets import QMessageBox
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle('Экспорт сеток')
        msg_box.setText(f'Что экспортировать?')
        msg_box.setInformativeText(f'Доступно категорий: {len(categories)}')
        btn_all = msg_box.addButton("Все категории", QMessageBox.YesRole)
        btn_current = msg_box.addButton("Только текущая", QMessageBox.NoRole)
        btn_cancel = msg_box.addButton("Отмена", QMessageBox.RejectRole)
        msg_box.setDefaultButton(btn_all)
        msg_box.exec_()
        
        clicked_button = msg_box.clickedButton()
        
        if clicked_button == btn_cancel:
            return
        
        export_all = (clicked_button == btn_all)
        
        if export_all:
            default_filename = "Все_сетки_турнира.docx"
            categories_to_export = categories
        else:
            if not self.current_category:
                QMessageBox.warning(self, "Ошибка", "Не выбрана категория для экспорта")
                return
            default_filename = f"{self.current_category}_сетка.docx"
            categories_to_export = [self.current_category]
        
        filename, _ = QFileDialog.getSaveFileName(
            self, 
            "Сохранить сетку в DOCX", 
            default_filename,
            "DOCX files (*.docx)"
        )
        
        if not filename:
            return

        try:
            # Создаем документ
            doc = Document()
            
            # Общий заголовок только для одной категории
            if not export_all:
                title = doc.add_heading(f'Сетка категории: {categories_to_export[0]}', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Экспортируем каждую категорию
            for idx, cat_name in enumerate(categories_to_export):
                if export_all and idx > 0:
                    doc.add_page_break()
                self._export_category_to_doc(doc, cat_name, tournament_data)
            
            # Сохраняем документ
            doc.save(filename)
            
            msg = f"Экспортировано категорий: {len(categories_to_export)}\nФайл: {filename}"
            QMessageBox.information(self, "Успех", msg)
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось экспортировать сетку:\n{str(e)}")
            import traceback
            print(traceback.format_exc())
    
    def _export_category_to_doc(self, doc, cat_name, tournament_data):
        """Экспортирует одну категорию в документ"""
        
        category = tournament_data.get('categories', {}).get(cat_name, {})
        
        cat_type = category.get('type', 'round_robin')
        
        
        # Сохраняем текущее состояние
        old_category = self.current_category
        
        # Временно обновляем данные для экспорта
        self.current_category = cat_name
        matches = category.get('matches', [])
        
        self._export_table_to_doc(doc, cat_name, tournament_data)
        doc.add_paragraph('')  # Пустая строка
        
        self._export_bracket_image_to_doc(doc, matches)
        
        # Восстанавливаем состояние
        if old_category:
            self.current_category = old_category
            self.update_bracket(old_category)
    
    def _export_table_to_doc(self, doc, cat_name, tournament_data):
        """Экспортирует табличное представление категории в документ"""
        # Временно обновляем таблицу для текущей категории
        self.update_round_robin_table(cat_name)
        
        table_widget = self.round_table
        if table_widget.rowCount() > 0 and table_widget.columnCount() > 0:
            # Создаем таблицу в документе
            doc_table = doc.add_table(rows=table_widget.rowCount() + 1, cols=table_widget.columnCount())
            doc_table.style = 'Light Grid Accent 1'
            
            # Заголовки
            for col in range(table_widget.columnCount()):
                header_item = table_widget.horizontalHeaderItem(col)
                if header_item:
                    cell = doc_table.rows[0].cells[col]
                    cell.text = header_item.text()
                    # Жирный шрифт для заголовков
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Данные таблицы
            for row in range(table_widget.rowCount()):
                for col in range(table_widget.columnCount()):
                    item = table_widget.item(row, col)
                    if item:
                        doc_table.rows[row + 1].cells[col].text = item.text()
                    else:
                        doc_table.rows[row + 1].cells[col].text = ""
        else:
            doc.add_paragraph('(Табличное представление пусто)')
    
    def _export_bracket_image_to_doc(self, doc, matches):
        """Экспортирует графическое представление в документ"""
        import tempfile
        import os
        
        from docx.shared import Inches
        
        # Создаем временный bracket widget для рендеринга
        temp_bracket = BracketWidget()
        temp_bracket.set_matches(matches)
        
        # Рендерим сцену напрямую в высоком разрешении для лучшего качества
        scene = temp_bracket.scene
        if scene and scene.items():
            # Получаем размеры сцены
            scene_rect = scene.sceneRect()
            if not scene_rect.isEmpty():
                # Высокое разрешение для качественного изображения (300 DPI)
                scale_factor = 3.0  # Увеличиваем разрешение в 3 раза
                width = int(scene_rect.width() * scale_factor)
                height = int(scene_rect.height() * scale_factor)
                
                # Создаем pixmap с высоким разрешением
                pixmap = QPixmap(width, height)
                pixmap.fill(Qt.white)
                
                # Рендерим сцену в pixmap с высоким качеством
                painter = QPainter(pixmap)
                painter.setRenderHint(QPainter.Antialiasing, True)
                painter.setRenderHint(QPainter.TextAntialiasing, True)
                painter.setRenderHint(QPainter.SmoothPixmapTransform, True)
                
                # Масштабируем для высокого разрешения
                painter.scale(scale_factor, scale_factor)
                painter.translate(-scene_rect.left(), -scene_rect.top())
                
                # Рендерим сцену
                scene.render(painter, scene_rect, scene_rect)
                painter.end()
                
                # Сохраняем во временный файл с высоким качеством
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    tmp_filename = tmp_file.name
                    # Сохраняем PNG (без потерь качества)
                    pixmap.save(tmp_filename, 'PNG')
                
                # Добавляем изображение в документ (уменьшено в 1.2 раза: 2 / 1.2 ≈ 1.67)
                try:
                    doc.add_picture(tmp_filename, width=Inches(2))
                    # Удаляем временный файл
                    os.unlink(tmp_filename)
                except Exception as e:
                    print(f"Ошибка при добавлении изображения: {e}")
                    doc.add_paragraph('(Не удалось добавить графическое представление)')
                    if os.path.exists(tmp_filename):
                        os.unlink(tmp_filename)
            else:
                doc.add_paragraph('(Графическое представление пусто)')
        else:
            doc.add_paragraph('(Графическое представление недоступно)')


class TournamentManager(QWidget):
    def __init__(self, parent=None, network_manager=None):
        super().__init__(parent)
        self.tournament_data = None
        self.current_match_index = 0
        self.network_manager = network_manager
        self.bracket_window = None
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        load_group = QGroupBox("Загрузка турнира")
        load_layout = QHBoxLayout(load_group)
        load_btn = QPushButton("Загрузить турнир из JSON")
        load_btn.clicked.connect(self.load_tournament)
        load_layout.addWidget(load_btn)
        self.tournament_label = QLabel("Турнир не загружен")
        load_layout.addWidget(self.tournament_label)
        layout.addWidget(load_group)

        self.info_group = QGroupBox("Информация о турнире")
        info_layout = QVBoxLayout(self.info_group)
        self.tournament_info = QTextEdit()
        self.tournament_info.setMaximumHeight(100)
        info_layout.addWidget(self.tournament_info)
        layout.addWidget(self.info_group)
        self.info_group.setVisible(False)

        self.management_group = QGroupBox("Управление категориями и участниками")
        management_layout = QHBoxLayout(self.management_group)

        left_layout = QVBoxLayout()
        left_layout.addWidget(QLabel("Категории"))

        self.categories_list = QListWidget()
        self.categories_list.currentItemChanged.connect(self.on_category_selected)
        left_layout.addWidget(self.categories_list)

        cat_btns = QHBoxLayout()
        add_cat_btn = QPushButton("Добавить")
        add_cat_btn.clicked.connect(self.add_category)
        cat_btns.addWidget(add_cat_btn)

        edit_cat_btn = QPushButton("Редактировать")
        edit_cat_btn.clicked.connect(self.edit_category)
        cat_btns.addWidget(edit_cat_btn)

        delete_cat_btn = QPushButton("Удалить")
        delete_cat_btn.clicked.connect(self.delete_category)
        cat_btns.addWidget(delete_cat_btn)

        left_layout.addLayout(cat_btns)
        management_layout.addLayout(left_layout, 1)

        right_layout = QVBoxLayout()
        right_layout.addWidget(QLabel("Участники категории"))

        self.participants_list = QListWidget()
        right_layout.addWidget(self.participants_list)

        part_btns = QHBoxLayout()
        add_part_btn = QPushButton("Добавить")
        add_part_btn.clicked.connect(self.add_wrestler)
        part_btns.addWidget(add_part_btn)

        remove_part_btn = QPushButton("Удалить")
        remove_part_btn.clicked.connect(self.remove_wrestler)
        part_btns.addWidget(remove_part_btn)

        move_part_btn = QPushButton("Переместить")
        move_part_btn.clicked.connect(self.move_wrestler)
        part_btns.addWidget(move_part_btn)

        color_part_btn = QPushButton("Цвет")
        color_part_btn.clicked.connect(self.change_wrestler_color)
        part_btns.addWidget(color_part_btn)

        right_layout.addLayout(part_btns)
        management_layout.addLayout(right_layout, 2)

        layout.addWidget(self.management_group)
        self.management_group.setVisible(False)

        self.matches_group = QGroupBox("Матчи турнира")
        matches_layout = QVBoxLayout(self.matches_group)

        category_layout = QHBoxLayout()
        category_layout.addWidget(QLabel("Категория:"))
        self.category_combo = QComboBox()
        self.category_combo.currentTextChanged.connect(self.on_category_combo_changed)
        category_layout.addWidget(self.category_combo)
        matches_layout.addLayout(category_layout)

        self.matches_list = QListWidget()
        self.matches_list.currentRowChanged.connect(self.select_match)
        matches_layout.addWidget(self.matches_list)

        current_match_layout = QHBoxLayout()
        self.prev_btn = QPushButton("Предыдущий")
        self.prev_btn.clicked.connect(self.previous_match)
        current_match_layout.addWidget(self.prev_btn)

        self.current_match_label = QLabel("Матч не выбран")
        current_match_layout.addWidget(self.current_match_label)

        self.next_btn = QPushButton("Следующий")
        self.next_btn.clicked.connect(self.next_match)
        current_match_layout.addWidget(self.next_btn)

        matches_layout.addLayout(current_match_layout)

        self.start_match_btn = QPushButton("Начать матч")
        self.start_match_btn.clicked.connect(self.start_current_match)
        matches_layout.addWidget(self.start_match_btn)

        self.round_robin_btn = QPushButton("Создать круговые сетки для всех категорий")
        self.round_robin_btn.clicked.connect(self.make_all_round_robin)
        matches_layout.addWidget(self.round_robin_btn)

        layout.addWidget(self.matches_group)
        self.matches_group.setVisible(False)

        open_bracket_btn = QPushButton("Открыть окно сетки")
        open_bracket_btn.clicked.connect(self.open_bracket_window)
        layout.addWidget(open_bracket_btn)

        open_mat_schedule_btn = QPushButton("Открыть расписание на ковре")
        open_mat_schedule_btn.clicked.connect(self.open_mat_schedule_window)
        layout.addWidget(open_mat_schedule_btn)

        transmit_btn = QPushButton("Передать категории на другой ПК")
        transmit_btn.clicked.connect(self.transmit_categories)
        layout.addWidget(transmit_btn)

        save_btn = QPushButton("Сохранить турнир")
        save_btn.clicked.connect(self.save_tournament)
        layout.addWidget(save_btn)

    def auto_load_into_control_panel(self, data):
        """Автоматическая загрузка данных матча в панель управления"""
        print(f"[DEBUG] Автозагрузка матча: {data}")
    
        # Прямой поиск панели управления для ковра 1
        cp = None
        
        # Вариант 1: Ищем среди всех виджетов
        for widget in QApplication.allWidgets():
            if hasattr(widget, 'mat_number') and widget.mat_number == 1:
                cp = widget
                break
        
        # Вариант 2: Если не нашли, ищем через главное окно
        if not cp:
            main_window = None
            for widget in QApplication.topLevelWidgets():
                # Ищем главное окно (EnhancedControlPanel)
                if hasattr(widget, 'open_control_panel_tab'):
                    main_window = widget
                    break
            
            if main_window:
                if hasattr(main_window, 'find_control_panel_by_mat'):
                    cp = main_window.find_control_panel_by_mat(1)
                elif hasattr(main_window, 'open_control_panel_tab'):
                    # Создаем панель управления, если ее нет
                    main_window.open_control_panel_tab(mat_number=1)
                    QApplication.processEvents()
                    # Снова ищем
                    for widget in QApplication.allWidgets():
                        if hasattr(widget, 'mat_number') and widget.mat_number == 1:
                            cp = widget
                            break
    
        if not cp:
            print("[DEBUG] Панель управления не найдена даже после открытия")
            return
        
        # Проверяем, есть ли уже активный незавершенный матч
        if hasattr(cp, 'current_match_w1') and hasattr(cp, 'current_match_w2'):
            if cp.current_match_w1 and cp.current_match_w2:
                # Проверяем, не завершен ли текущий матч
                if hasattr(cp, 'current_match_id') and cp.current_match_id:
                    # Ищем матч в категории
                    category = self.tournament_data.get('categories', {}).get(data.get("category", ""), {})
                    matches = category.get('matches', [])
                    for m in matches:
                        if m.get('id') == cp.current_match_id:
                            # Если матч не завершен, не заменяем его
                            if not m.get('completed', False):
                                print(f"[DEBUG] Пропускаем автозагрузку: уже есть активный матч {cp.current_match_w1} vs {cp.current_match_w2}")
                                return
                            break
                else:
                    # Если есть активные борцы, но нет match_id, тоже не заменяем
                    print(f"[DEBUG] Пропускаем автозагрузку: уже есть активный матч {cp.current_match_w1} vs {cp.current_match_w2}")
                    return
    
        # Получаем данные о клубах
        w1_name = data.get("w1", "")
        w2_name = data.get("w2", "")
        
        w1_club = ""
        w2_club = ""
        
        if self.tournament_data:
            # Ищем клуб в данных турнира
            for participant in self.tournament_data.get('participants', []):
                if participant.get('name') == w1_name:
                    w1_club = participant.get('club', '')
                    break
            
            for participant in self.tournament_data.get('participants', []):
                if participant.get('name') == w2_name:
                    w2_club = participant.get('club', '')
                    break
    
        w1 = {
            "name": w1_name,
            "club": w1_club
        }
        
        w2 = {
            "name": w2_name,
            "club": w2_club
        }
    
        print(f"[DEBUG] Установка борцов: {w1['name']} vs {w2['name']}")
    
        # Устанавливаем борцов
        cp.set_match_competitors(w1, w2)
    
        if hasattr(cp, 'set_current_match_info'):
            cp.set_current_match_info(
                data.get("category", ""),
                w1_name,
                w2_name,
                data.get("match_id")
            )
    
        # Если матч уже завершен, загружаем счет
        if data.get("completed", False):
            print(f"[DEBUG] Загружаем счет: {data.get('score1', 0)} - {data.get('score2', 0)}")
            # Проверяем, кто был wrestler1 в данных матча
            if data.get("w1") == w1_name:
                cp.red.points = data.get("score1", 0)
                cp.blue.points = data.get("score2", 0)
            else:
                cp.red.points = data.get("score2", 0)
                cp.blue.points = data.get("score1", 0)
            
            # Обновляем отображение
            if hasattr(cp, 'update_display'):
                cp.update_display()
    
        # Немедленная отправка на табло
        cp.send_scoreboard_update()
        
        # Дополнительная отправка через 100 мс для надежности
        QTimer.singleShot(100, cp.send_scoreboard_update)
        
        print(f"[DEBUG] Матч успешно загружен в панель управления")

    def open_bracket_window(self):
        if self.bracket_window is None:
            self.bracket_window = BracketWindow(self, self.tournament_data)
            self.bracket_window.match_autoload.connect(self.auto_load_into_control_panel)

        self.bracket_window.show()
        self.bracket_window.raise_()
        self.bracket_window.activateWindow()

        current_cat = self.category_combo.currentText()
        if current_cat:
            self.bracket_window.update_bracket(current_cat)
            # Автозагрузка только если нет активного матча
            # autoload_match теперь сам проверяет наличие активного матча
            self.bracket_window.autoload_match()

    def open_mat_schedule_window(self):
        """Открывает отдельное окно расписания на ковре прямо из менеджера турнира."""
        if not self.tournament_data:
            QMessageBox.warning(self, "Внимание", "Сначала загрузите турнир")
            return
        if not hasattr(self, 'mat_schedule_window') or self.mat_schedule_window is None:
            self.mat_schedule_window = MatScheduleWindow(self.tournament_data, self, self.network_manager)
            self.mat_schedule_window.destroyed.connect(lambda: setattr(self, 'mat_schedule_window', None))
        self.mat_schedule_window.show()
        self.mat_schedule_window.raise_()
        self.mat_schedule_window.activateWindow()

    def on_category_combo_changed(self, cat):
        self.update_matches_list(cat)
        if self.bracket_window:
            self.bracket_window.update_bracket(cat)

    def load_tournament(self):
        filename, _ = QFileDialog.getOpenFileName(self, "Открыть турнир", "", "JSON files (*.json)")
        if filename:
            self.load_tournament_from_file(filename)

    def load_tournament_from_file(self, filename):
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                self.tournament_data = json.load(f)

            self.tournament_label.setText(f"Загружен: {self.tournament_data.get('name', 'Без имени')}")
            self.update_tournament_info()
            self.update_categories_lists()
            self.generate_tournament_schedule()

            self.info_group.setVisible(True)
            self.management_group.setVisible(True)
            self.matches_group.setVisible(True)

            if self.bracket_window:
                self.bracket_window.tournament_data = self.tournament_data
                current_cat = self.category_combo.currentText()
                if current_cat:
                    self.bracket_window.update_bracket(current_cat)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить: {e}")

    def update_tournament_info(self):
        if not self.tournament_data:
            return
        info = f"""
        Название: {self.tournament_data.get('name', '')}
        Дата: {self.tournament_data.get('date', '')}
        Место: {self.tournament_data.get('location', '')}
        Участников: {len(self.tournament_data.get('participants', []))}
        Категорий: {len(self.tournament_data.get('categories', {}))}
        """
        self.tournament_info.setPlainText(info.strip())

    def update_categories_lists(self):
        self.categories_list.clear()
        self.category_combo.clear()

        if not self.tournament_data:
            return

        for cat in self.tournament_data['categories'].keys():
            self.categories_list.addItem(cat)
            self.category_combo.addItem(cat)

    def on_category_selected(self, current, previous):
        if not current:
            self.participants_list.clear()
            return

        cat = current.text()
        self.participants_list.clear()
        participants = self.tournament_data['categories'][cat]['participants']
        for p in participants:
            self.participants_list.addItem(p['name'])

    def add_category(self):
        dialog = CategoryEditDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            name = dialog.get_name()
            gender = dialog.get_gender()
            age = dialog.get_age()
            weight = dialog.get_weight()

            if name in self.tournament_data['categories']:
                QMessageBox.warning(self, "Ошибка", "Категория уже существует")
                return

            self.tournament_data['categories'][name] = {
                'gender': gender,
                'age': age,
                'weight': weight,
                'participants': [],
                'matches': [],
                'type': 'elimination'
            }
            self.update_categories_lists()
            self.generate_tournament_schedule()

    def edit_category(self):
        item = self.categories_list.currentItem()
        if not item:
            return
        old_name = item.text()
        dialog = CategoryEditDialog(self, old_name,
                                   self.tournament_data['categories'][old_name])
        if dialog.exec_() == QDialog.Accepted:
            new_name = dialog.get_name()
            if new_name != old_name and new_name in self.tournament_data['categories']:
                QMessageBox.warning(self, "Ошибка", "Категория уже существует")
                return

            data = self.tournament_data['categories'].pop(old_name)
            data['gender'] = dialog.get_gender()
            data['age'] = dialog.get_age()
            data['weight'] = dialog.get_weight()
            self.tournament_data['categories'][new_name] = data

            for p in self.tournament_data['participants']:
                if p.get('category') == old_name:
                    p['category'] = new_name

            self.update_categories_lists()
            self.generate_tournament_schedule()

    def delete_category(self):
        item = self.categories_list.currentItem()
        if not item:
            return
        cat = item.text()
        reply = QMessageBox.question(self, "Удалить", f"Удалить категорию {cat}?", QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            del self.tournament_data['categories'][cat]
            self.tournament_data['participants'] = [
                p for p in self.tournament_data['participants'] if p.get('category') != cat
            ]
            self.update_categories_lists()
            self.generate_tournament_schedule()

    def add_wrestler(self):
        cat_item = self.categories_list.currentItem()
        if not cat_item:
            QMessageBox.warning(self, "Ошибка", "Выберите категорию")
            return
        cat = cat_item.text()

        dialog = AddWrestlerDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            new = dialog.get_wrestler()

            for p in self.tournament_data['categories'][cat]['participants']:
                if p['name'] == new['name']:
                    QMessageBox.warning(self, "Ошибка", "Такой участник уже есть в категории")
                    return

            self.tournament_data['categories'][cat]['participants'].append({
                'name': new['name'],
                'club': new.get('club', ''),
                'color': new.get('color', '')
            })

            self.tournament_data.setdefault('participants', [])
            self.tournament_data['participants'].append({
                'name': new['name'],
                'club': new.get('club', ''),
                'category': cat,
                'color': new.get('color', '')
            })

            self.on_category_selected(cat_item, None)
            self.generate_tournament_schedule()

            if self.bracket_window:
                self.bracket_window.update_bracket(cat)

    def remove_wrestler(self):
        cat_item = self.categories_list.currentItem()
        part_item = self.participants_list.currentItem()
        if not cat_item or not part_item:
            return
        cat = cat_item.text()
        name = part_item.text()

        reply = QMessageBox.question(self, "Удалить", f"Удалить {name}?", QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.tournament_data['categories'][cat]['participants'] = [
                p for p in self.tournament_data['categories'][cat]['participants'] if p['name'] != name
            ]
            self.tournament_data['participants'] = [
                p for p in self.tournament_data['participants'] if p['name'] != name
            ]
            self.regenerate_bracket(cat)
            self.on_category_selected(cat_item, None)
            self.generate_tournament_schedule()
            # Синхронизируем изменения
            self._sync_tournament_changes()

    def move_wrestler(self):
        cat_item = self.categories_list.currentItem()
        part_item = self.participants_list.currentItem()
        if not cat_item or not part_item:
            return
        current_cat = cat_item.text()
        name = part_item.text()

        dialog = MoveWrestlerDialog(self.tournament_data['categories'], current_cat, self)
        if dialog.exec_() == QDialog.Accepted:
            target_cat = dialog.get_target()
            if target_cat == current_cat:
                return

            for i, p in enumerate(self.tournament_data['categories'][current_cat]['participants']):
                if p['name'] == name:
                    moved = self.tournament_data['categories'][current_cat]['participants'].pop(i)
                    break

            for p in self.tournament_data['participants']:
                if p['name'] == name:
                    p['category'] = target_cat
                    break

            self.tournament_data['categories'][target_cat]['participants'].append(moved)
            self.regenerate_bracket(current_cat)
            self.regenerate_bracket(target_cat)
            self.on_category_selected(cat_item, None)
            self.generate_tournament_schedule()

    def change_wrestler_color(self):
        """Изменение цвета участника с обновлением расписаний."""
        cat_item = self.categories_list.currentItem()
        part_item = self.participants_list.currentItem()
        if not cat_item or not part_item:
            QMessageBox.warning(self, "Ошибка", "Выберите участника")
            return
        cat = cat_item.text()
        name = part_item.text()
        color = QColorDialog.getColor()
        if not color.isValid():
            return
        color_hex = color.name()

        # Обновляем в категории
        for p in self.tournament_data['categories'].get(cat, {}).get('participants', []):
            if p.get('name') == name:
                p['color'] = color_hex
                break

        # Обновляем в глобальном списке участников
        for p in self.tournament_data.get('participants', []):
            if p.get('name') == name:
                p['color'] = color_hex
                break

        # Перегенерируем расписание, чтобы цвета применились
        self.generate_tournament_schedule()
        if self.bracket_window:
            self.bracket_window.update_bracket(cat)

        # Обновляем отображение выбранного элемента
        part_item.setBackground(QBrush(QColor(color_hex)))

    def regenerate_bracket(self, cat):
        wrestlers = self.tournament_data['categories'][cat]['participants']
        new_bracket = create_bracket(wrestlers, cat)
        self.tournament_data['categories'][cat]['matches'] = new_bracket['matches']
        self.tournament_data['categories'][cat]['type'] = new_bracket['type']

        if self.bracket_window:
            self.bracket_window.update_bracket(cat)
        
        # Синхронизируем изменения
        self._sync_tournament_changes()
    
    def _sync_tournament_changes(self):
        """Синхронизирует изменения турнира через schedule_sync."""
        schedule_sync = self._get_schedule_sync()
        if schedule_sync and self.tournament_data:
            try:
                schedule_sync.push_schedule(self.tournament_data)
                print(f"[SYNC] Изменения в турнире синхронизированы")
            except Exception as e:
                print(f"[ERROR] Ошибка синхронизации изменений турнира: {e}")
    
    def _get_schedule_sync(self):
        """Получает schedule_sync из родительского окна."""
        parent = self.parent()
        while parent:
            if hasattr(parent, 'schedule_sync_service'):
                return parent.schedule_sync_service
            try:
                parent = parent.parent()
            except (AttributeError, RuntimeError):
                break
        # Если не нашли через родителя, ищем через QApplication
        from PyQt5.QtWidgets import QApplication
        for window in QApplication.topLevelWidgets():
            if hasattr(window, 'schedule_sync_service'):
                return window.schedule_sync_service
        return None

    def make_all_round_robin(self):
        if not self.tournament_data or 'categories' not in self.tournament_data:
            QMessageBox.warning(self, "Ошибка", "Турнир не загружен")
            return

        for cat, data in self.tournament_data['categories'].items():
            wrestlers = data.get('participants', [])
            bracket = create_bracket(wrestlers, cat, bracket_type='round_robin')
            data['matches'] = bracket.get('matches', [])
            data['type'] = bracket.get('type', 'round_robin')

        self.generate_tournament_schedule()

        current_cat = self.category_combo.currentText()
        if current_cat:
            self.update_matches_list(current_cat)
            if self.bracket_window:
                self.bracket_window.update_bracket(current_cat)

        QMessageBox.information(self, "Готово", "Для всех категорий созданы круговые сетки.")

    def generate_tournament_schedule(self):
        if not self.tournament_data:
            return
        try:
            settings = get_settings()
            # Перезагружаем настройки перед генерацией
            settings.load_settings()
            n_mats = settings.get("tournament", "number_of_mats", 2)
            print(f"[DEBUG tournament_manager.generate_tournament_schedule] Прочитано n_mats={n_mats} (тип: {type(n_mats).__name__})")
            if n_mats < 1:
                n_mats = 2  # Минимум 2 ковра
                settings.set("tournament", "number_of_mats", n_mats)
                print(f"[WARNING] Количество ковров было меньше 1, установлено значение {n_mats}")
            schedule = generate_schedule(self.tournament_data, start_time="10:00", match_duration=8, n_mats=n_mats)
            self.tournament_data["schedule"] = schedule
            print(f"[INFO] Расписание сгенерировано для {n_mats} ковров")
            main_window = self.window()
            if hasattr(main_window, 'update_schedule_tab'):
                main_window.update_schedule_tab()
            if hasattr(self, 'mat_schedule_window') and self.mat_schedule_window:
                self.mat_schedule_window.update_data(self.tournament_data)
            # Синхронизируем изменения
            self._sync_tournament_changes()
        except Exception as e:
            print(f"Ошибка расписания: {e}")
            import traceback
            traceback.print_exc()

    def update_matches_list(self, cat):
        self.matches_list.clear()
        if not cat or not self.tournament_data:
            return
        matches = self.tournament_data['categories'][cat].get('matches', [])
        
        # Сортируем матчи по раундам для правильного отображения алгоритма round-robin
        # Сначала по раунду, потом по порядку в раунде
        matches_sorted = sorted(matches, key=lambda m: (
            m.get('round', 1),  # Сначала по раунду
            m.get('id', '')     # Потом по ID для стабильной сортировки
        ))

        for m in matches_sorted:
            round_num = m.get('round', 1)
            text = f"Раунд {round_num}: {m['wrestler1']} vs {m['wrestler2']}"
            if m.get('winner'):
                text += f" → {m['winner']}"
            self.matches_list.addItem(text)

    def select_match(self, row):
        if row < 0:
            return
        cat = self.category_combo.currentText()
        matches = self.tournament_data['categories'][cat]['matches']
        if row < len(matches):
            match = matches[row]
            self.current_match_label.setText(f"{match['wrestler1']} vs {match['wrestler2']}")

    def previous_match(self):
        current = self.matches_list.currentRow()
        if current > 0:
            self.matches_list.setCurrentRow(current - 1)

    def next_match(self):
        current = self.matches_list.currentRow()
        max_row = self.matches_list.count() - 1
        if current < max_row:
            self.matches_list.setCurrentRow(current + 1)

    def start_current_match(self):
        cat = self.category_combo.currentText()
        row = self.matches_list.currentRow()
        if row < 0 or not cat:
            return
        matches = self.tournament_data['categories'][cat]['matches']
        # Сортируем матчи по раундам для правильного порядка
        matches_sorted = sorted(matches, key=lambda m: (
            m.get('round', 1),
            m.get('id', '')
        ))
        if row >= len(matches_sorted):
            return
        match = matches_sorted[row]

        main_window = self.window()
        if main_window and hasattr(main_window, 'open_control_panel_tab'):
            main_window.open_control_panel_tab(mat_number=1)
            cp = main_window.find_control_panel_by_mat(1)
            if cp:
                cp.set_match_competitors(
                    {'name': match['wrestler1'], 'club': '', 'category': cat},
                    {'name': match['wrestler2'], 'club': '', 'category': cat}
                )
                if hasattr(cp, 'set_current_match_info'):
                    cp.set_current_match_info(
                        cat,
                        match.get('wrestler1'),
                        match.get('wrestler2'),
                        match.get('id')
                    )
                cp.send_scoreboard_update()

    def save_tournament(self):
        if not self.tournament_data:
            QMessageBox.warning(self, "Ошибка", "Нет данных для сохранения")
            return
        filename, _ = QFileDialog.getSaveFileName(self, "Сохранить турнир", "", "JSON files (*.json)")
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(self.tournament_data, f, ensure_ascii=False, indent=2, default=str)
                self.load_tournament_from_file(filename)
                QMessageBox.information(self, "Успех", "Турнир сохранен и перезагружен")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", str(e))

    def transmit_categories(self):
        if not self.tournament_data or 'categories' not in self.tournament_data:
            QMessageBox.warning(self, "Ошибка", "Нет категорий для передачи")
            return

        host = '192.168.1.100'
        port = 12345

        try:
            categories_data = json.dumps(self.tournament_data['categories'], ensure_ascii=False)
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.connect((host, port))
            sock.sendall(categories_data.encode('utf-8'))
            sock.close()
            QMessageBox.information(self, "Успех", "Категории переданы на другой ПК")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось передать категории: {str(e)}")


class NetworkServer(QTcpServer):
    def __init__(self, tournament_manager):
        super().__init__()
        self.tournament_manager = tournament_manager
        if self.listen(QHostAddress.Any, 12345):
            print("Сервер запущен на порту 12345")
        else:
            print("Ошибка запуска сервера")

    def incomingConnection(self, socketDescriptor):
        socket = QTcpSocket()
        socket.setSocketDescriptor(socketDescriptor)
        socket.readyRead.connect(self.read_data)
        socket.disconnected.connect(socket.deleteLater)

    def read_data(self):
        socket = self.sender()
        data = socket.readAll().data().decode('utf-8')
        try:
            categories = json.loads(data)
            self.tournament_manager.tournament_data['categories'] = categories
            self.tournament_manager.update_categories_lists()
            self.tournament_manager.generate_tournament_schedule()
            if self.tournament_manager.bracket_window:
                current_cat = self.tournament_manager.category_combo.currentText()
                if current_cat:
                    self.tournament_manager.bracket_window.update_bracket(current_cat)
            QMessageBox.information(self.tournament_manager, "Успех", "Категории получены от другого ПК")
        except Exception as e:
            QMessageBox.critical(self.tournament_manager, "Ошибка", f"Ошибка обработки данных: {str(e)}")